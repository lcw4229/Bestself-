#!/usr/bin/env python3
"""Generate International Sport Business Plan as a Word document."""

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
CHART_PATH = os.path.join(OUTPUT_DIR, "budget_chart.png")
DOC_PATH = os.path.join(OUTPUT_DIR, "International_Sport_Business_Plan.docx")

BRAND_ORANGE = RGBColor(0xE8, 0x6C, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MEDIUM_GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_paragraph(doc, text, bold=False, size=11,
                             color=DARK_GRAY, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             space_after=6, space_before=0, first_line_indent=None):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(first_line_indent)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = 'Times New Roman'
    return p


def add_heading_styled(doc, text, level=1):
    if level == 1:
        add_formatted_paragraph(doc, text, bold=True, size=16,
                                color=BRAND_ORANGE, space_before=18, space_after=8)
    elif level == 2:
        add_formatted_paragraph(doc, text, bold=True, size=13,
                                color=DARK_GRAY, space_before=12, space_after=6)


def add_body(doc, text, indent=False):
    return add_formatted_paragraph(doc, text, size=12, space_after=6,
                                    first_line_indent=0.5 if indent else None)


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.left_indent = Inches(0.5)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_GRAY
    run.font.name = 'Times New Roman'


def generate_budget_chart():
    categories = [
        'Facility Lease &\nRenovation',
        'Equipment &\nTechnology',
        'Staff Salaries\n& Training',
        'Digital Platform\nDevelopment',
        'Marketing &\nPromotion',
        'Operations &\nAdministration',
        'Contingency\nFund'
    ]
    year1 = [180000, 95000, 220000, 150000, 120000, 85000, 50000]
    year2 = [140000, 60000, 310000, 80000, 160000, 110000, 45000]
    year3 = [120000, 45000, 420000, 60000, 200000, 140000, 40000]

    revenue_years = ['Year 1', 'Year 2', 'Year 3']
    revenue = [320000, 780000, 1450000]
    expenses = [sum(year1), sum(year2), sum(year3)]

    fig, axes = plt.subplots(1, 2, figsize=(14, 5.5))

    x = np.arange(len(categories))
    width = 0.25
    colors = ['#E86C00', '#2E86AB', '#A23B72']

    bars1 = axes[0].bar(x - width, [v/1000 for v in year1], width, label='Year 1', color=colors[0], edgecolor='white', linewidth=0.5)
    bars2 = axes[0].bar(x, [v/1000 for v in year2], width, label='Year 2', color=colors[1], edgecolor='white', linewidth=0.5)
    bars3 = axes[0].bar(x + width, [v/1000 for v in year3], width, label='Year 3', color=colors[2], edgecolor='white', linewidth=0.5)

    axes[0].set_ylabel('Amount (USD Thousands)', fontsize=10, fontweight='bold')
    axes[0].set_title('Projected Expenses by Category', fontsize=12, fontweight='bold', pad=12)
    axes[0].set_xticks(x)
    axes[0].set_xticklabels(categories, fontsize=7.5)
    axes[0].legend(fontsize=9)
    axes[0].yaxis.set_major_formatter(ticker.FuncFormatter(lambda v, _: f'${v:.0f}K'))
    axes[0].spines['top'].set_visible(False)
    axes[0].spines['right'].set_visible(False)
    axes[0].grid(axis='y', alpha=0.3)

    x2 = np.arange(len(revenue_years))
    width2 = 0.35
    axes[1].bar(x2 - width2/2, [v/1000 for v in revenue], width2, label='Revenue', color='#2ECC71', edgecolor='white', linewidth=0.5)
    axes[1].bar(x2 + width2/2, [v/1000 for v in expenses], width2, label='Total Expenses', color='#E74C3C', edgecolor='white', linewidth=0.5)

    axes[1].set_ylabel('Amount (USD Thousands)', fontsize=10, fontweight='bold')
    axes[1].set_title('Revenue vs. Expenses Projection', fontsize=12, fontweight='bold', pad=12)
    axes[1].set_xticks(x2)
    axes[1].set_xticklabels(revenue_years, fontsize=10)
    axes[1].legend(fontsize=9)
    axes[1].yaxis.set_major_formatter(ticker.FuncFormatter(lambda v, _: f'${v:.0f}K'))
    axes[1].spines['top'].set_visible(False)
    axes[1].spines['right'].set_visible(False)
    axes[1].grid(axis='y', alpha=0.3)

    for bar_group in [bars1, bars2, bars3]:
        for bar in bar_group:
            h = bar.get_height()
            if h > 30:
                axes[0].annotate(f'${h:.0f}K', xy=(bar.get_x() + bar.get_width()/2, h),
                               xytext=(0, 3), textcoords='offset points',
                               ha='center', va='bottom', fontsize=6.5)

    plt.tight_layout(pad=2.0)
    plt.savefig(CHART_PATH, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()


def build_document():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = DARK_GRAY

    # ── COVER PAGE ──
    for _ in range(6):
        doc.add_paragraph()

    add_formatted_paragraph(doc, 'HoopRise Basketball Academy', bold=True, size=28,
                             color=BRAND_ORANGE, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_formatted_paragraph(doc, 'International Sport Business Plan', bold=True, size=18,
                             color=DARK_GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 40)
    run.font.color.rgb = BRAND_ORANGE
    run.font.size = Pt(12)

    add_formatted_paragraph(doc, 'Expanding Youth Basketball Development in India',
                             size=14, color=MEDIUM_GRAY,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)
    add_formatted_paragraph(doc, 'Landon Worrich',
                             size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_formatted_paragraph(doc, 'SPMT 337: International Sport Business',
                             size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_formatted_paragraph(doc, 'June 2026',
                             size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    doc.add_page_break()

    # ── TABLE OF CONTENTS ──
    add_formatted_paragraph(doc, 'Table of Contents', bold=True, size=16,
                             color=BRAND_ORANGE, space_after=12)
    toc_items = [
        ('Market Research Summary', '1'),
        ('The Concept', '2'),
        ('Mission Statement', '3'),
        ('Competitive Position', '3'),
        ('Cultural Considerations', '5'),
        ('Fundraising/Financing', '6'),
        ('Promotional Strategies', '7'),
        ('Anticipated Budget', '8'),
        ('Strategic Partnerships', '9'),
        ('References', '10'),
    ]
    for title, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=1)
        run = p.add_run(title)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = DARK_GRAY
        run2 = p.add_run(f'\t{page}')
        run2.font.size = Pt(12)
        run2.font.name = 'Times New Roman'
        run2.font.color.rgb = MEDIUM_GRAY

    doc.add_page_break()

    # ── MARKET RESEARCH SUMMARY ──
    add_heading_styled(doc, 'Market Research Summary')

    add_heading_styled(doc, 'Key Demographics and Market Size', level=2)
    add_body(doc,
        'We looked at a bunch of different countries before settling on India. Southeast Asia '
        'came up, so did parts of Africa and a few spots in Eastern Europe. But once we actually '
        'sat down with the data on India, it wasn\'t even close. 1.44 billion people live there. '
        'Around 65% of them are under 35, and the median age is 28.4 (United Nations, 2024). '
        'That\'s not just a big market. That\'s an absurdly young country full of people at the '
        'exact age where they\'d pick up a new sport if someone gave them the chance.',
        indent=True)

    add_body(doc,
        'India\'s sports industry is worth about $2.7 billion today. Analysts think it\'ll hit '
        '$10 billion by 2030 (KPMG, 2023). And here\'s the number that really got us going: the '
        'NBA says it already has over 300 million fans in India. That makes India their '
        'second-largest market outside the States. 300 million people who care about basketball, '
        'and almost nowhere for them to actually learn how to play it properly. That disconnect '
        'is the whole reason HoopRise exists.',
        indent=True)

    add_heading_styled(doc, 'Consumer Behavior Trends', level=2)
    add_body(doc,
        'Young Indians today aren\'t like their parents. They\'re on Instagram at midnight '
        'watching Steph Curry highlights. They care about fitness. A lot of this traces back to '
        'how insanely cheap mobile data got in India, like $0.17 per GB cheap (Deloitte, 2024). '
        'Everyone has a smartphone now, and that means everyone\'s consuming sports content all '
        'the time.',
        indent=True)

    add_body(doc,
        'On the parent side, India\'s middle class has blown up to over 400 million people. '
        'These families want their kids doing more than just sitting in tutoring classes after '
        'school. They\'re open to sports in a way that previous generations just weren\'t. And if '
        'anyone doubts whether Indians will get behind a well-packaged sport product, just look '
        'at what the IPL did for cricket. It printed money and built a cultural obsession at the '
        'same time. There\'s no reason basketball can\'t follow a similar path.',
        indent=True)

    add_heading_styled(doc, 'Opportunities and Challenges', level=2)
    add_body(doc,
        'The timing\'s good. The government\'s Khelo India program has dumped over $350 million '
        'into sports development. The NBA opened an academy in India in 2017, which means smart '
        'people with deep pockets already see what we see. And with 35% of the population living '
        'in cities now, there are concentrated areas of demand that a physical academy can '
        'actually serve.',
        indent=True)

    add_body(doc,
        'That said, we\'d be lying if we pretended there weren\'t real obstacles. Cricket isn\'t '
        'just popular in India. It\'s basically a religion. Basketball is fighting for attention '
        'against something that deeply embedded. Good indoor courts are tough to find outside the '
        'biggest metros, and even middle-class Indian families are pretty price-conscious, so the '
        'training can\'t cost a fortune. India also has 28 states with different regulations and '
        'cultural norms, so you can\'t just copy-paste one strategy across the whole country. '
        'We\'ve factored all of that into the plan, but it\'d be dishonest to pretend those '
        'challenges don\'t exist.',
        indent=True)

    doc.add_page_break()

    # ── THE CONCEPT ──
    add_heading_styled(doc, 'The Concept')
    add_body(doc,
        'HoopRise is pretty simple at its core. We\'re building real basketball academies in '
        'three major Indian cities (Mumbai, Delhi NCR, Bangalore), and we\'re pairing those with '
        'a mobile app that brings our training to anyone with a phone and a Wi-Fi connection. '
        'Physical locations for the kids who can get there. A digital product for the millions '
        'who can\'t.',
        indent=True)

    add_body(doc,
        'The academies will have full-size courts, weight rooms, sports science labs with motion '
        'capture, and film rooms for game breakdown. Coaches will be people who\'ve played or '
        'coached internationally, paired with local assistants who know the area and can actually '
        'relate to the kids. Training splits into four tiers: Beginner, Intermediate, Advanced, '
        'Elite. An 8-year-old who\'s never held a ball and a 21-year-old trying to go pro both '
        'need a place to train. We want to be that place for both of them.',
        indent=True)

    add_body(doc,
        'The app is where things get exciting from a business standpoint. AI skill assessments, '
        'drill videos in multiple Indian languages, stat tracking, virtual coaching. If you\'re a '
        'kid in, say, Jaipur or Indore right now and you want real basketball coaching? Good '
        'luck. There\'s basically nothing. That\'s the gap we\'re filling. The app makes quality '
        'training available to anyone with a phone, while the physical academies give the full '
        'experience to players in our three launch cities. Nobody else in India is doing both.',
        indent=True)

    doc.add_page_break()

    # ── MISSION STATEMENT ──
    add_heading_styled(doc, 'Mission Statement')
    add_body(doc,
        'Where you grow up shouldn\'t decide if you get a good coach. How much money your family '
        'has shouldn\'t either. That\'s the whole idea behind HoopRise. We want to build India\'s '
        'best basketball training program and then make it available to as many young players as '
        'we can, not just the ones in wealthy families or the right zip code. And it\'s not only '
        'about basketball. The discipline, the accountability, learning to work with people who '
        'are different from you, all of that matters just as much as getting your shot right.',
        indent=True)

    add_body(doc,
        'We\'ve got three goals. One: become the name people think of when Indian basketball '
        'training comes up, within five years. Two: build a real pipeline of players who end up '
        'in college programs, national teams, or pro leagues. Three: get to a million app users '
        'in three years. Big numbers. We know. But India\'s young population is massive, the '
        'interest is there, and nobody\'s really going after this market with a serious plan yet. '
        'So we\'re going for it.',
        indent=True)

    # ── COMPETITIVE POSITION ──
    add_heading_styled(doc, 'Competitive Position')

    add_heading_styled(doc, 'Target Market Profile', level=2)
    add_body(doc,
        'Main target is 8-to-22-year-olds from middle-class and upper-middle-class families in '
        'Indian cities. Households making $10,000 to $50,000 a year, give or take. Parents who '
        'value education but are starting to realize sports can be part of the picture too. The '
        'kids themselves grew up on the internet. They watch the NBA. A lot of them already want '
        'to play. They just don\'t have a good place to do it.',
        indent=True)

    add_body(doc,
        'Beyond that, there\'s a secondary market: college athletes who want better coaching, '
        'companies interested in team-building stuff, and adults in their twenties and thirties '
        'who play pickup on weekends. About 120 million young people across India\'s top 20 '
        'cities fit our core profile. We\'re starting in Mumbai (21 million metro area), Delhi '
        'NCR (32 million), and Bangalore (13 million) because those three have the best overlap '
        'of middle-class population, basketball interest, and available real estate for facilities.',
        indent=True)

    add_heading_styled(doc, 'Key Competitors', level=2)

    table = doc.add_table(rows=5, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    headers = ['Competitor', 'Strengths', 'Weaknesses', 'Market Position']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        run.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'E86C00')

    competitors = [
        ['NBA Academy India',
         'Huge brand; world-class coaches; players can get scouted by the NBA directly',
         'Takes like 24 kids. Total. One single location. If you\'re not already elite, forget it.',
         'Niche/Elite'],
        ['Stepanova Basketball\nAcademy',
         'Been around for a while; cheap; grassroots focused',
         'Tiny facilities; no app or tech; some coaches are great, others not so much',
         'Local/Grassroots'],
        ['Local Municipal\nPrograms',
         'Free or close to it; government funded; available in lots of cities',
         'Courts in rough shape; coaches who often have zero formal training; no curriculum',
         'Mass/Basic'],
        ['International Apps\n(HomeCourt, etc.)',
         'Solid tech; AI feedback; big global user base',
         'No one on the ground in India; English only; designed for Americans, not Indians',
         'Digital/Global'],
    ]

    for row_idx, comp in enumerate(competitors, 1):
        for col_idx, val in enumerate(comp):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = DARK_GRAY
            if row_idx % 2 == 0:
                set_cell_shading(cell, 'FFF3E6')

    add_formatted_paragraph(doc, '', size=6, space_after=6)

    add_heading_styled(doc, 'Competitive Differentiation', level=2)
    add_body(doc,
        'Look at that table and the gap is pretty obvious. The NBA Academy? Amazing, but it '
        'serves a couple dozen kids. Municipal programs? Open to everyone, but the quality is '
        'rough. Apps like HomeCourt? Cool tech, zero connection to India. Nobody\'s doing both '
        'the physical and digital side in this market. That\'s the lane HoopRise is taking.',
        indent=True)

    add_body(doc,
        'Premium enough to actually develop real players, but priced so a normal family can '
        'swing it. Specifically, here\'s what separates us:',
        indent=True)

    add_bullet(doc, 'Physical plus Digital: Train at the academy, keep working through the app '
                    'at home. The learning doesn\'t end when you walk out the door.')
    add_bullet(doc, 'Personalized Plans: AI tracks what each player\'s good at and what they\'re '
                    'struggling with, then adjusts. No one\'s running the same generic drills.')
    add_bullet(doc, 'Scholarships Built In: Financial aid was part of the model from the start. '
                    'If a kid\'s talented, money shouldn\'t be what stops them.')
    add_bullet(doc, 'Multilingual: App launches in Hindi, English, Tamil, Kannada, and Marathi. '
                    'Players should be able to train in a language they actually think in.')
    add_bullet(doc, 'A Real Path Forward: Direct connections to international college programs '
                    'and pro leagues. Not vague promises, actual relationships.')

    doc.add_page_break()

    # ── CULTURAL CONSIDERATIONS ──
    add_heading_styled(doc, 'Cultural Considerations')

    add_heading_styled(doc, 'Family-Centric Decision Making', level=2)
    add_body(doc,
        'In India, the kid doesn\'t decide to join a basketball program. The family does. '
        'Sometimes it\'s mom and dad. Sometimes it\'s grandparents. Sometimes an uncle has '
        'opinions. That means our marketing can\'t just be highlight reels and cool branding '
        'aimed at teenagers. We have to answer the questions parents are actually asking: Will '
        'this help my kid\'s discipline? Will it teach them time management? Could this lead to '
        'a college scholarship?',
        indent=True)

    add_body(doc,
        'So every piece of marketing hits those points. We\'ll run open-house days where '
        'families can walk through the facility, talk to coaches, ask whatever they want. Parents '
        'get quarterly progress reports. The whole signup process is designed to make them feel '
        'like they\'re part of this, not just the ones paying for it. If the parents aren\'t '
        'bought in, the kid isn\'t showing up. That\'s just how it works there.',
        indent=True)

    add_heading_styled(doc, 'Education-First Culture', level=2)
    add_body(doc,
        'School comes first. Always. If an Indian parent has to pick between a tutor and '
        'basketball practice, basketball loses ten out of ten times. We\'re not fighting that '
        'because we\'d lose. Instead, we built the schedule around it. Practice runs outside '
        'school hours. During board exam season, which hits in March-April and again in '
        'October-November, we cut back sessions and let families pick flexible times. Every '
        'academy has a study room where kids can do homework before or after they train, and '
        'that detail matters more than it might seem. It tells parents we\'re not pulling their '
        'kid away from academics.',
        indent=True)

    add_body(doc,
        'The scholarship angle is big, too. We\'ll show families real stories of basketball '
        'players who got into international universities because of the sport. When a parent sees '
        'that basketball could actually help their kid\'s education rather than hurt it, the '
        'conversation changes completely.',
        indent=True)

    add_heading_styled(doc, 'Language and Regional Diversity', level=2)
    add_body(doc,
        'India has 22 official languages. Hundreds of dialects on top of that. Running everything '
        'in English only would be a huge mistake. Coaches at each academy will work bilingually: '
        'English and Hindi in Delhi, English and Marathi in Mumbai, English and Kannada in '
        'Bangalore. The app ships in five languages at launch, with more coming based on what '
        'people ask for. And we\'re localizing marketing, too. Not just translating it, but '
        'actually using cultural references that mean something in each city. A campaign that hits '
        'in Delhi might fall totally flat in Bangalore if you don\'t adjust the context.',
        indent=True)

    add_heading_styled(doc, 'Religious and Festival Sensitivity', level=2)
    add_body(doc,
        'Diwali. Eid. Christmas. Holi. Pongal. Navratri. The list goes on. India\'s festival '
        'calendar is packed, and these aren\'t just days off. They\'re deeply important to '
        'families. We\'re building our yearly schedule with those dates baked in from the start. '
        'Training adjusts around major holidays. During festival seasons, we\'ll host themed '
        'events that bring together kids from all different backgrounds. Making the academy a '
        'place where everyone feels welcome isn\'t optional. In a country this diverse, it\'s the '
        'only way to keep people coming back long-term.',
        indent=True)

    add_heading_styled(doc, 'Gender Considerations', level=2)
    add_body(doc,
        'Women\'s basketball is growing in urban India, but there\'s still real hesitation in '
        'some families about girls training alongside boys, especially once they\'re teenagers. '
        'We\'re offering dedicated girls-only sessions with female coaches. That one change '
        'removes the biggest barrier for families who\'d otherwise just say no.',
        indent=True)

    add_body(doc,
        'We\'ll also make sure female players are front and center in our marketing. Not shoved '
        'into a sidebar, but actually featured as a core part of the brand. And if you look at '
        'it purely from a business perspective, ignoring the women\'s side means voluntarily '
        'cutting our addressable market in half. That just doesn\'t make sense.',
        indent=True)

    doc.add_page_break()

    # ── FUNDRAISING ──
    add_heading_styled(doc, 'Fundraising/Financing')

    add_body(doc,
        'Total cost to launch and survive year one: roughly $900,000. That\'s facilities, staff, '
        'app development, marketing, and a buffer for things we didn\'t predict. We\'re pulling '
        'that from four different places so we\'re not dependent on any single source.',
        indent=True)

    add_heading_styled(doc, 'Angel Investors and Venture Capital', level=2)
    add_body(doc,
        'Sports-tech investing in India has gotten a lot more active. Dream Sports, Blume '
        'Ventures, Sequoia India, they\'re all looking for plays in this space. We\'re targeting '
        'a $400,000 seed round from angels with backgrounds in sports, education, or youth '
        'development. The pitch is pretty straightforward: India has the world\'s biggest young '
        'population, basketball interest is spiking, and nobody\'s built a scalable training '
        'model that mixes physical and digital. That\'s the gap. We\'re filling it.',
        indent=True)

    add_heading_styled(doc, 'Government Grants and Initiatives', level=2)
    add_body(doc,
        'India\'s Khelo India program and the Sports Authority of India hand out grants to '
        'organizations doing grassroots sports work. If we get accredited as an official training '
        'center, that could mean $100,000 to $150,000 in grant money plus cheap access to '
        'government sports facilities. Maharashtra, Delhi, and Karnataka all have their own '
        'state-level sports funds too, so we\'re applying at every level we can.',
        indent=True)

    add_heading_styled(doc, 'Corporate Sponsorships', level=2)
    add_body(doc,
        'Most people outside India don\'t know this, but Indian companies above a certain profit '
        'level are legally required to spend 2% of their net profits on corporate social '
        'responsibility. Youth sports qualifies. That means we can walk into a meeting with Tata, '
        'Reliance, or Infosys and present a sponsorship deal where they\'re spending money they '
        'have to spend anyway, but on something that actually gives them brand exposure with young '
        'consumers. We\'re putting together sponsorship tiers from $25,000 to $200,000, with '
        'visibility at academies, in the app, and at events.',
        indent=True)

    add_heading_styled(doc, 'Crowdfunding', level=2)
    add_body(doc,
        'We\'ll run a campaign on Ketto (India\'s biggest crowdfunding platform) and Kickstarter. '
        'Target is $50,000 to $75,000, but the money\'s almost secondary. What crowdfunding '
        'really does is build a base of early supporters who are personally invested before we\'ve '
        'even opened the doors. Backers get discounted memberships, merch, and early app access. '
        'It also forces us to tell our story publicly, and stories about giving kids access to '
        'sports they couldn\'t otherwise afford tend to travel fast online.',
        indent=True)

    doc.add_page_break()

    # ── PROMOTIONAL STRATEGIES ──
    add_heading_styled(doc, 'Promotional Strategies')

    add_heading_styled(doc, 'Digital Marketing Campaigns', level=2)
    add_body(doc,
        '800 million internet users. Digital\'s obviously the main channel. Targeted ads on '
        'Instagram, YouTube, and Facebook, heavy on video. Training clips, player progress '
        'stories, behind-the-scenes academy stuff. Placing YouTube pre-roll ads on NBA content '
        'is kind of a no-brainer since that\'s exactly where our audience already hangs out.',
        indent=True)

    add_body(doc,
        'We\'ll also invest in SEO and Google Ads for search terms like "basketball training '
        'India" and "basketball academy near me." And we\'ll put out weekly blog posts and '
        'tutorial videos. The idea isn\'t just advertising, it\'s making HoopRise the name that '
        'keeps popping up whenever anyone in India searches for anything related to basketball '
        'development.',
        indent=True)

    add_heading_styled(doc, 'Influencer and Ambassador Partnerships', level=2)
    add_body(doc,
        'We want actual people attached to this brand. Satnam Singh was the first Indian ever '
        'drafted by an NBA team. Prachi Tehlan captained the national women\'s team. Those are '
        'the kinds of names we\'re going after. In each city, we\'ll also work with local '
        'micro-influencers, people with 50K to 200K followers who have real pull in their '
        'communities, not just big numbers. We want content collabs and event appearances that '
        'feel genuine, not stiff endorsement deals that everyone scrolls past.',
        indent=True)

    add_heading_styled(doc, 'School and University Outreach', level=2)
    add_body(doc,
        'Our customers sit in classrooms five days a week. Getting into schools is a huge deal. '
        'Free intro clinics at partner schools, equipment donations for schools starting '
        'basketball programs, inter-school tournament sponsorships. On the college side, we want '
        'to be the training partner that university basketball teams turn to for group packages '
        'and performance analytics through the app. It\'s about being present in the places where '
        'young people already are, instead of trying to pull them somewhere new.',
        indent=True)

    add_heading_styled(doc, 'Event Sponsorships and Community Engagement', level=2)
    add_body(doc,
        'Online ads can only do so much. People need to experience the product. So we\'re doing '
        '3-on-3 street tournaments. Basketball camps during school breaks. Watch parties for the '
        'NBA Finals and All-Star Weekend. And once a year, a big event called the "HoopRise '
        'Rising Stars Challenge" that gets press, drives social media buzz, and gives the brand '
        'a signature moment. Long term, we want HoopRise to feel like the center of basketball '
        'culture in India. Not just a training academy, but the place where the basketball '
        'community actually gathers.',
        indent=True)

    doc.add_page_break()

    # ── ANTICIPATED BUDGET ──
    add_heading_styled(doc, 'Anticipated Budget')

    add_body(doc,
        'Three-year financial picture. Year 1 we spend a lot and don\'t make it all back. Year 2 '
        'we grow into profitability. Year 3 we\'re making money.',
        indent=True)

    budget_table = doc.add_table(rows=9, cols=4)
    budget_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    budget_table.style = 'Table Grid'

    budget_headers = ['Category', 'Year 1 (USD)', 'Year 2 (USD)', 'Year 3 (USD)']
    for i, h in enumerate(budget_headers):
        cell = budget_table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        run.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'E86C00')

    budget_data = [
        ['Facility Lease & Renovation', '$180,000', '$140,000', '$120,000'],
        ['Equipment & Technology', '$95,000', '$60,000', '$45,000'],
        ['Staff Salaries & Training', '$220,000', '$310,000', '$420,000'],
        ['Digital Platform Development', '$150,000', '$80,000', '$60,000'],
        ['Marketing & Promotion', '$120,000', '$160,000', '$200,000'],
        ['Operations & Administration', '$85,000', '$110,000', '$140,000'],
        ['Contingency Fund', '$50,000', '$45,000', '$40,000'],
        ['Total Expenses', '$900,000', '$905,000', '$1,025,000'],
    ]

    for row_idx, row_data in enumerate(budget_data, 1):
        for col_idx, val in enumerate(row_data):
            cell = budget_table.rows[row_idx].cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            if row_idx == 8:
                run.bold = True
                run.font.color.rgb = WHITE
                set_cell_shading(cell, '333333')
            else:
                run.font.color.rgb = DARK_GRAY
                if col_idx > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if row_idx % 2 == 0:
                    set_cell_shading(cell, 'FFF3E6')

    add_formatted_paragraph(doc, '', size=8, space_after=8)

    add_heading_styled(doc, 'Projected Revenue', level=2)

    rev_table = doc.add_table(rows=5, cols=4)
    rev_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rev_table.style = 'Table Grid'

    rev_headers = ['Revenue Stream', 'Year 1 (USD)', 'Year 2 (USD)', 'Year 3 (USD)']
    for i, h in enumerate(rev_headers):
        cell = rev_table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        run.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '2E86AB')

    rev_data = [
        ['Academy Memberships', '$180,000', '$420,000', '$720,000'],
        ['Digital Platform Subscriptions', '$40,000', '$150,000', '$380,000'],
        ['Corporate Programs & Events', '$60,000', '$130,000', '$230,000'],
        ['Sponsorship & Partnerships', '$40,000', '$80,000', '$120,000'],
    ]

    for row_idx, row_data in enumerate(rev_data, 1):
        for col_idx, val in enumerate(row_data):
            cell = rev_table.rows[row_idx].cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = DARK_GRAY
            if col_idx > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if row_idx % 2 == 0:
                set_cell_shading(cell, 'E6F4FA')

    add_formatted_paragraph(doc, '', size=6, space_after=4)

    add_body(doc,
        'Year 1 loses money. We expected that. You can\'t open three academies and build an app '
        'without burning through cash. But revenue ramps up quick. Somewhere around month 18, we '
        'break even. By Year 3, we\'re pulling in $1.45 million against about $1 million in '
        'costs. That\'s solid margin, and it gets better as the app scales because digital users '
        'don\'t require more facilities or proportionally more coaches.',
        indent=True)

    add_formatted_paragraph(doc, '', size=4, space_after=2)
    doc.add_picture(CHART_PATH, width=Inches(6.2))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ── STRATEGIC PARTNERSHIPS ──
    add_heading_styled(doc, 'Strategic Partnerships')

    add_body(doc,
        'Trying to do everything ourselves would be slow, expensive, and kind of dumb when '
        'there are organizations out there that already have what we need. The smarter move is '
        'to partner up, trade value, and get to market faster. Five partnerships matter most.',
        indent=True)

    add_heading_styled(doc, 'NBA India', level=2)
    add_body(doc,
        'This is the big one. The NBA\'s already spending money in India through the Academy, '
        'Jr. NBA programs, and broadcast deals. Getting recognized as an official NBA-affiliated '
        'training center would give us their coaching resources, their credibility, and access to '
        'their scouting pipeline. For a kid training at HoopRise, that means a real ladder: '
        'perform well, get noticed, maybe get a shot at the Academy or catch a scout\'s eye. No '
        'other training program in India can dangle that kind of carrot.',
        indent=True)

    add_heading_styled(doc, 'Nike India / Adidas India', level=2)
    add_body(doc,
        'Sportswear partnership gives us better prices on gear, co-branded marketing, and '
        'possibly scholarship funding. Nike\'s our first pick because of their basketball roots, '
        'but Adidas works too. For the brand, they get their products on the next generation of '
        'Indian basketball players in a market that\'s only going to get bigger. They spend '
        'millions trying to build that kind of ground-level brand loyalty through other channels. '
        'We\'re offering it built in.',
        indent=True)

    add_heading_styled(doc, 'Basketball Federation of India (BFI)', level=2)
    add_body(doc,
        'BFI accreditation gives us institutional credibility you can\'t buy. Our players become '
        'eligible for state and national team selection. We get into BFI tournaments and coaching '
        'certification programs. For the BFI, we\'re doing the grassroots scouting and '
        'development they don\'t have the infrastructure to do themselves. Both sides get '
        'something they need.',
        indent=True)

    add_heading_styled(doc, 'Indian School and University Networks', level=2)
    add_body(doc,
        'Delhi Public School, Amity International, Ryan International. These school networks '
        'have campuses everywhere. Partnering with them means after-school programs on their '
        'facilities, discounted enrollment for their students, and inter-school tournaments we '
        'sponsor. Universities like Christ (Bangalore), Shiv Nadar (Delhi), and NMIMS (Mumbai) '
        'give us access to the 18-22 age bracket and potential internship pipelines for sports '
        'management students. It puts us right where our customers spend most of their day.',
        indent=True)

    add_heading_styled(doc, 'Technology Partners', level=2)
    add_body(doc,
        'Good apps are expensive to build. Great apps even more so. Partnering with Indian IT '
        'companies like Infosys or Wipro for development, using AWS India\'s startup credits for '
        'hosting, and integrating tools from sports tech companies like ShotTracker keeps our '
        'costs manageable without sacrificing quality. When you\'re a startup trying to compete '
        'against polished international apps, you need the tech to be good. These partnerships '
        'make that possible without blowing through our budget.',
        indent=True)

    doc.add_page_break()

    # ── REFERENCES ──
    add_heading_styled(doc, 'References')

    references = [
        'Deloitte. (2024). India\'s digital sports consumption: Trends and opportunities in a '
        'mobile-first market. Deloitte Touche Tohmatsu India LLP.',

        'KPMG. (2023). The business of sports: A comprehensive analysis of India\'s evolving '
        'sports industry landscape. KPMG India.',

        'National Basketball Association. (2024). NBA global outreach report: Expanding '
        'basketball culture in emerging markets. NBA Communications.',

        'Sports Authority of India. (2023). Khelo India annual report 2022-2023: Building a '
        'sporting nation through grassroots development. Ministry of Youth Affairs and Sports, '
        'Government of India.',

        'United Nations Department of Economic and Social Affairs. (2024). World population '
        'prospects 2024: India demographic profile summary. United Nations Publications.',
    ]

    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        run = p.add_run(ref)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = DARK_GRAY

    doc.save(DOC_PATH)
    print(f"Document saved: {DOC_PATH}")


if __name__ == '__main__':
    generate_budget_chart()
    print(f"Chart saved: {CHART_PATH}")
    build_document()
